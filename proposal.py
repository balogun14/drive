from docx import Document
from docx.shared import Inches

# Create the document
doc = Document()
doc.add_heading('Project Proposal: School Management System', 0)

doc.add_paragraph('Prepared by: [Insert Your Business Name Here]')
doc.add_paragraph('Company Logo: [Insert Your Company Logo Here]')
doc.add_paragraph('Date: [Insert Date]')

doc.add_heading('Introduction', level=1)
doc.add_paragraph(
    'We propose the development and deployment of a comprehensive School Management System (SMS) '
    'that automates and simplifies academic, administrative, and financial operations within schools. '
    'This solution empowers administrators, teachers, students, and parents to communicate effectively '
    'and manage responsibilities with ease.'
)

doc.add_heading('Project Goals', level=1)
goals = [
    'Streamline student enrollment and academic tracking',
    'Simplify grade management and performance reporting',
    'Enable secure and seamless communication among stakeholders',
    'Automate financial transactions such as fee payments and receipts',
    'Offer real-time data insights for better decision-making'
]
for goal in goals:
    doc.add_paragraph(f'• {goal}', style='List Bullet')

doc.add_heading('Key Features', level=1)

doc.add_heading('User Management', level=2)
doc.add_paragraph('- Dedicated profiles for students, teachers, parents, and administrators')
doc.add_paragraph('- Role-based access with permissions tailored to each user type')

doc.add_heading('Academic Management', level=2)
doc.add_paragraph('- Class creation and subject allocation')
doc.add_paragraph('- Timetable and attendance management')
doc.add_paragraph('- Assignment, exam, and assessment tracking')

doc.add_heading('Assessment & Grading', level=2)
doc.add_paragraph('- Online assignment and exam creation')
doc.add_paragraph('- Submission, evaluation, and grading workflows')
doc.add_paragraph('- Automated GPA and report card generation')

doc.add_heading('Fee Management', level=2)
doc.add_paragraph('- Fee schedule generation')
doc.add_paragraph('- Online and offline payment support')
doc.add_paragraph('- Receipt issuance and transaction history')

doc.add_heading('Communication', level=2)
doc.add_paragraph('- Email and SMS notifications for key events')
doc.add_paragraph('- Parent-teacher communication tools')
doc.add_paragraph('- Alerts for grades, payments, announcements')

doc.add_heading('Reporting & Analytics', level=2)
doc.add_paragraph('- Downloadable reports (PDF/Excel) on academic performance, finances, and system usage')
doc.add_paragraph('- Real-time dashboards for administrators')

doc.add_heading('Security Features', level=1)
security_features = [
    'Role-Based Access Control (RBAC)',
    'Encrypted user data and credentials',
    'Audit trails to monitor changes and activities',
    'Secure integration with payment gateways'
]
for item in security_features:
    doc.add_paragraph(f'• {item}', style='List Bullet')

doc.add_heading('Integration Capabilities', level=1)
integration_features = [
    'Payment Gateway (for fee transactions)',
    'Email & SMS Notification Services',
    'Cloud-based File Storage',
    'Export functionalities (PDF/Excel)'
]
for item in integration_features:
    doc.add_paragraph(f'• {item}', style='List Bullet')

doc.add_heading('System Architecture Overview', level=1)
doc.add_paragraph(
    '[Web Interface] <--> [Application Layer] <--> [Database / File Storage / Cache]\n\n'
    'Main Services:\n'
    '- Authentication Service\n'
    '- User Service\n'
    '- Student/Teacher Services\n'
    '- Grade & Report Services\n'
    '- Fee & Payment Services'
)

doc.add_heading('User Workflows', level=1)

doc.add_heading('1. Student Enrollment', level=2)
doc.add_paragraph('From application submission to class allocation and welcome kit issuance.')

doc.add_heading('2. Teacher Assignment Grading', level=2)
doc.add_paragraph('Teachers create assignments, collect submissions, grade them, and notify students.')

doc.add_heading('3. Parent Fee Payment', level=2)
doc.add_paragraph('Parents log in, review dues, make payments, and receive digital receipts.')

doc.add_heading('4. Admin Reporting', level=2)
doc.add_paragraph('Admin users generate academic and financial reports, download or view them online.')

doc.add_heading('Benefits to Your School', level=1)
benefits = [
    'Efficiency: Automates repetitive tasks, reducing administrative burden',
    'Transparency: Clear visibility into academic and financial records',
    'Engagement: Fosters stronger communication between parents and teachers',
    'Scalability: Designed to grow with your school'
]
for benefit in benefits:
    doc.add_paragraph(f'• {benefit}', style='List Bullet')

doc.add_heading('Next Steps', level=1)
doc.add_paragraph(
    'If you\'re ready to transform your school\'s operations with our School Management System, '
    'we’d love to schedule a discovery meeting.\n\n'
    '✅ Please insert your business name and logo above before sending this document to clients.'
)

doc.add_heading('Contact Information', level=1)
doc.add_paragraph('[Your Name]')
doc.add_paragraph('[Your Email Address]')
doc.add_paragraph('[Your Phone Number]')
doc.add_paragraph('[Your Website or Portfolio Link]')

# Save the document
file_path = "/mnt/data/School_Management_System_Proposal.docx"
doc.save(file_path)

file_path
